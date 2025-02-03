import streamlit as st
from groq import Groq
from PyPDF2 import PdfReader
from docx import Document
from tiktoken import get_encoding
import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64
import os
from typing import List, Dict, Any
import concurrent.futures
from datetime import datetime

# Initialize Groq client (use environment variable in production)
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

class DocumentProcessor:
    @staticmethod
    def extract_text_from_pdf(file) -> str:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text

    @staticmethod
    def extract_text_from_docx(file) -> str:
        doc = Document(file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    @staticmethod
    def preprocess_text(text: str) -> str:
        return " ".join(text.replace("\n", " ").replace("\r", " ").split())

class TextChunker:
    def __init__(self):
        self.encoding = get_encoding("cl100k_base")

    def split_into_chunks(self, text: str, token_limit: int = 5500) -> List[str]:
        words = text.split()
        chunks = []
        current_chunk = []
        current_tokens = 0

        for word in words:
            word_tokens = len(self.encoding.encode(word + " "))
            if current_tokens + word_tokens > token_limit:
                chunks.append(" ".join(current_chunk))
                current_chunk = [word]
                current_tokens = word_tokens
            else:
                current_chunk.append(word)
                current_tokens += word_tokens

        if current_chunk:
            chunks.append(" ".join(current_chunk))
        return chunks

class LegalAnalyzer:
    def __init__(self, client: Groq):
        self.client = client
        self.key_clauses = {
            "non_compete": {
                "keywords": ["non-compete", "noncompete", "competitive activities", "competition restriction"],
                "importance": "HIGH"
            },
            "limitation_of_liability": {
                "keywords": ["limitation of liability", "limited liability", "liability cap", "maximum liability"],
                "importance": "HIGH"
            },
            "indemnification": {
                "keywords": ["indemnify", "indemnification", "hold harmless", "indemnity"],
                "importance": "HIGH"
            },
            "termination": {
                "keywords": ["termination", "terminate", "cancellation", "right to terminate"],
                "importance": "HIGH"
            },
            "force_majeure": {
                "keywords": ["force majeure", "act of god", "unforeseen circumstances"],
                "importance": "MEDIUM"
            },
            "confidentiality": {
                "keywords": ["confidential", "confidentiality", "non-disclosure", "proprietary information"],
                "importance": "HIGH"
            },
            "governing_law": {
                "keywords": ["governing law", "jurisdiction", "venue", "applicable law"],
                "importance": "MEDIUM"
            },
            "amendment": {
                "keywords": ["amendment", "modification", "modify", "changes to agreement"],
                "importance": "MEDIUM"
            }
        }

    def summarize_chunk(self, chunk: str) -> str:
        try:
            response = self.client.chat.completions.create(
                messages=[{
                    "role": "user",
                    "content": f"Summarize the following legal document in a concise manner: {chunk}"
                }],
                model="llama-3.1-8b-instant",
                stream=False
            )
            if response and response.choices:
                return response.choices[0].message.content
            return "Error: Empty response from API"
        except Exception as e:
            return f"Error: {str(e)}"

    def analyze_document_structure(self, text: str) -> Dict[str, Any]:
        sections = [
            "definitions", "parties", "term", "payment", "confidentiality",
            "intellectual property", "termination", "governing law",
            "indemnification", "warranties", "force majeure"
        ]
        
        structure = {}
        for section in sections:
            pos = text.lower().find(section)
            if pos != -1:
                context_start = max(0, pos - 100)
                context_end = min(len(text), pos + 500)
                structure[section] = {
                    "position": pos,
                    "context": text[context_start:context_end].strip(),
                    "has_subsections": bool(text[pos:pos+1000].count("\n\n") > 2)
                }
        
        return structure

    def analyze_risks(self, text: str, summary: str) -> Dict[str, Any]:
        risk_categories = {
            "financial_risks": {
                "keywords": ["payment", "penalty", "fee", "damages", "costs", "expenses"],
                "severity_words": ["substantial", "significant", "material"]
            },
            "legal_risks": {
                "keywords": ["liability", "breach", "violation", "lawsuit", "litigation"],
                "severity_words": ["serious", "severe", "critical"]
            },
            "operational_risks": {
                "keywords": ["delay", "failure", "interruption", "termination"],
                "severity_words": ["immediate", "substantial", "significant"]
            },
            "compliance_risks": {
                "keywords": ["regulation", "requirement", "law", "policy", "standard"],
                "severity_words": ["mandatory", "required", "essential"]
            }
        }

        results = {}
        for category, data in risk_categories.items():
            category_risks = []
            for keyword in data["keywords"]:
                positions = self._find_all_positions(text.lower(), keyword)
                for pos in positions:
                    context = self._extract_context(text, pos, keyword)
                    severity = self._assess_severity(context, data["severity_words"])
                    category_risks.append({
                        "keyword": keyword,
                        "context": context,
                        "severity": severity,
                        "position": pos
                    })
            
            if category_risks:
                results[category] = {
                    "risks": category_risks,
                    "total_count": len(category_risks),
                    "severity_distribution": self._calculate_severity_distribution(category_risks)
                }

        return results

    def detect_key_clauses(self, text: str) -> Dict[str, Any]:
        text_lower = text.lower()
        detected_clauses = {}

        for clause_type, clause_info in self.key_clauses.items():
            clause_instances = []
            for keyword in clause_info["keywords"]:
                positions = self._find_all_positions(text_lower, keyword.lower())
                for pos in positions:
                    context = self._extract_context(text, pos, keyword)
                    clause_instances.append({
                        "keyword": keyword,
                        "context": context,
                        "position": pos,
                        "importance": clause_info["importance"]
                    })
            
            if clause_instances:
                detected_clauses[clause_type] = {
                    "instances": clause_instances,
                    "count": len(clause_instances),
                    "importance": clause_info["importance"]
                }

        return detected_clauses

    def _find_all_positions(self, text: str, keyword: str) -> List[int]:
        positions = []
        start = 0
        while True:
            pos = text.find(keyword, start)
            if pos == -1:
                break
            positions.append(pos)
            start = pos + 1
        return positions

    def _extract_context(self, text: str, position: int, keyword: str, window: int = 200) -> str:
        start = max(0, position - window)
        end = min(len(text), position + len(keyword) + window)
        return text[start:end].strip()

    def _assess_severity(self, context: str, severity_words: List[str]) -> str:
        context_lower = context.lower()
        if any(word in context_lower for word in severity_words):
            return "HIGH"
        return "MEDIUM" if "may" in context_lower or "should" in context_lower else "LOW"

    def _calculate_severity_distribution(self, risks: List[Dict[str, Any]]) -> Dict[str, int]:
        distribution = {"HIGH": 0, "MEDIUM": 0, "LOW": 0}
        for risk in risks:
            distribution[risk["severity"]] += 1
        return distribution

class Visualizer:
    @staticmethod
    def create_risk_visualizations(analyzed_risks: Dict[str, Any]) -> plt.Figure:
        fig = plt.figure(figsize=(15, 10))
        
        # Risk Category Distribution
        plt.subplot(2, 2, 1)
        categories = list(analyzed_risks.keys())
        counts = [data["total_count"] for data in analyzed_risks.values()]
        plt.bar(categories, counts)
        plt.title("Risk Distribution by Category")
        plt.xticks(rotation=45)
        
        # Severity Distribution
        plt.subplot(2, 2, 2)
        severity_data = {
            "HIGH": sum(data["severity_distribution"]["HIGH"] for data in analyzed_risks.values()),
            "MEDIUM": sum(data["severity_distribution"]["MEDIUM"] for data in analyzed_risks.values()),
            "LOW": sum(data["severity_distribution"]["LOW"] for data in analyzed_risks.values())
        }
        plt.pie(severity_data.values(), labels=severity_data.keys(), autopct='%1.1f%%')
        plt.title("Risk Severity Distribution")
        
        # Risk Severity Heatmap
        plt.subplot(2, 2, 3)
        heatmap_data = [[data["severity_distribution"][sev] for sev in ["HIGH", "MEDIUM", "LOW"]] 
                        for data in analyzed_risks.values()]
        sns.heatmap(heatmap_data, 
                    xticklabels=["HIGH", "MEDIUM", "LOW"],
                    yticklabels=categories,
                    annot=True,
                    fmt="d",
                    cmap="YlOrRd")
        plt.title("Risk Severity Heatmap")
        
        plt.tight_layout()
        return fig

class ReportGenerator:
    @staticmethod
    def generate_analysis_report(document_text: str, summary: str, structure: Dict[str, Any], 
                               analyzed_risks: Dict[str, Any]) -> Document:
        doc = Document()
        
        # Title and metadata
        doc.add_heading('Legal Document Analysis Report', level=1)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=2)
        doc.add_paragraph(summary)
        
        # Document Structure Analysis
        doc.add_heading('Document Structure', level=2)
        for section, details in structure.items():
            doc.add_heading(section.title(), level=3)
            doc.add_paragraph(f"Location: {details['position']}")
            doc.add_paragraph(f"Context: {details['context']}")
        
        # Risk Analysis
        doc.add_heading('Risk Analysis', level=2)
        for category, data in analyzed_risks.items():
            doc.add_heading(category.replace('_', ' ').title(), level=3)
            doc.add_paragraph(f"Total Risks: {data['total_count']}")
            
            # Severity Distribution
            doc.add_paragraph("Severity Distribution:")
            for severity, count in data["severity_distribution"].items():
                doc.add_paragraph(f"- {severity}: {count}", style='List Bullet')
            
            # Detailed Risks
            doc.add_paragraph("Detailed Risks:")
            for risk in data["risks"]:
                doc.add_paragraph(f"Keyword: {risk['keyword']}", style='List Bullet')
                doc.add_paragraph(f"Severity: {risk['severity']}")
                doc.add_paragraph(f"Context: {risk['context']}")
        
        return doc

def display_legal_analysis_page():
    st.title("Advanced Legal Document Analysis")
    
    # Initialize components
    processor = DocumentProcessor()
    chunker = TextChunker()
    analyzer = LegalAnalyzer(client)
    visualizer = Visualizer()
    
    uploaded_file = st.file_uploader("Upload Legal Document", type=["pdf", "docx"])
    
    if uploaded_file:
        # Extract and process text
        if uploaded_file.name.endswith('.pdf'):
            text = processor.extract_text_from_pdf(uploaded_file)
        else:
            text = processor.extract_text_from_docx(uploaded_file)
        
        processed_text = processor.preprocess_text(text)
        
        # Create tabs for different analyses
        tabs = st.tabs(["Document Text", "Summary", "Structure Analysis", "Risk Analysis", "Report"])
        
        with tabs[0]:
            st.subheader("Document Text")
            st.text_area("Extracted Text", processed_text, height=300)
        
        with tabs[1]:
            st.subheader("Document Summary")
            chunks = chunker.split_into_chunks(processed_text)
            with st.spinner("Generating summary..."):
                summaries = [analyzer.summarize_chunk(chunk) for chunk in chunks]
                summary = " ".join(summaries)
                st.write(summary)
        
        with tabs[2]:
            st.subheader("Document Structure Analysis")
            structure = analyzer.analyze_document_structure(processed_text)
            for section, details in structure.items():
                with st.expander(f"{section.title()} Section"):
                    st.write(f"Position in document: {details['position']}")
                    st.write("Context:")
                    st.text(details['context'])
        
        with tabs[3]:
            st.subheader("Risk and Key Clause Analysis")
            analyzed_risks = analyzer.analyze_risks(processed_text, summary)
            key_clauses = analyzer.detect_key_clauses(processed_text)
            
            # Create two columns
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Risk Analysis")
                # Display visualizations
                if analyzed_risks:
                    fig = visualizer.create_risk_visualizations(analyzed_risks)
                    st.pyplot(fig)
                    
                    # Display detailed risk analysis
                    for category, data in analyzed_risks.items():
                        with st.expander(f"{category.replace('_', ' ').title()} ({data['total_count']} risks)"):
                            st.write(f"**Severity Distribution:**")
                            for severity, count in data["severity_distribution"].items():
                                st.write(f"- {severity}: {count}")
                            
                            st.write("\n**Detailed Risks:**")
                            for risk in data["risks"]:
                                st.markdown(f"""
                                - **Keyword:** {risk['keyword']}
                                - **Severity:** {risk['severity']}
                                - **Context:** "{risk['context']}"
                                ---
                                """)
            
            with col2:
                st.subheader("Key Clauses")
                if key_clauses:
                    for clause_type, clause_data in key_clauses.items():
                        with st.expander(f"{clause_type.replace('_', ' ').title()} ({clause_data['count']} instances)"):
                            st.write(f"**Importance:** {clause_data['importance']}")
                            st.write("\n**Detected Instances:**")
                            for instance in clause_data['instances']:
                                st.markdown(f"""
                                - **Keyword Found:** {instance['keyword']}
                                - **Context:** "{instance['context']}"
                                ---
                                """)
                else:
                    st.warning("No key clauses detected in the document.")
        
        with tabs[4]:
            st.subheader("Analysis Report")
            report_doc = ReportGenerator.generate_analysis_report(
                processed_text, summary, structure, analyzed_risks)
            
            # Save report
            report_path = "analysis_report.docx"
            report_doc.save(report_path)
            
            with open(report_path, "rb") as f:
                st.download_button(
                    "Download Full Analysis Report",
                    f,
                    file_name="legal_document_analysis_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
#add the code tab 6

if __name__ == "__main__":
    display_legal_analysis_page()